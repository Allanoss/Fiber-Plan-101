from django.contrib import messages
from calendar import c
from ssl import AlertDescription
from turtle import color
from django.http import HttpRequest
from django.shortcuts import redirect
import openrouteservice as ors
import math
import gmplot
import json
from django.urls import reverse
import googlemaps
import os
import gmaps.geojson_geometries
from ipywidgets.embed import embed_minimal_html
from requests import request

from plotter.models import Coordinate

Client1 = ors.Client(key='5b3ce3597851110001cf624820c44dd22f3c4a128773998371409a6f') #ors api key 1
Client2 = ors.Client(key='5b3ce3597851110001cf6248e6b29a12715344b7a18a42c037cceb60') #ors api key 2
apikey = 'AIzaSyA_xtmQprzuxWdhVy0JxjDPMA0loHExtx8'                                   #GoogleMaps api key

#invert coordinates
def Gmaps_distance(coordinates): 
    gmaps=googlemaps.Client(key=apikey)
    MinTot=999999999
    MaxTot=0

    Start = list()
    Altstart=list()

    for i in coordinates:
        Tdistance=0
        Gstart=i
        for y in coordinates:
           end = y
           route=gmaps.directions((Gstart[1], Gstart[0]),(end[1], end[0]), mode="driving", units='metric')
           distance=route[0]['legs'][0]['distance']['value']  
           Tdistance+=distance
        
        if MinTot>Tdistance:
              MinTot=Tdistance
              Start=i

        #Altstart to be used to create redundancy for network security
        if Tdistance>MaxTot:
              MaxTot=Tdistance
              Altstart=i 
    return Start, Altstart 

def ors_route(coordinates):
    route=Client1.directions(coordinates=coordinates, 
                             profile='foot-walking',
                             optimized= False,
                             preference='shortest',
                             continue_straight=True,                        
                             format='geojson')
    return route
     
def ors_distance(coordinates):
    route=Client2.directions(coordinates=coordinates, 
                             profile='foot-walking',
                             optimized= False,
                             preference='shortest',
                             continue_straight=True,                        
                             format='geojson')
             
    return route['features'][0]['properties']['segments'][0]['distance']    

def Start_Distance(points, connect_point=None):
    Start, Altstart= Gmaps_distance(points)
    for i in points: 
        Tdistance=0    
        now=[Start,i]
        try:
            distance=ors_distance(now)
        except:
           def Coorderror (request,id):
                messages.error(request, message='Check coordinates and try again.')
                   
        Tdistance+=distance
                   

    if connect_point:
        now=[Start, connect_point]
        try:
            distance=ors_distance(now)
        except:
           def Coorderror (request,id):
                messages.error(request, 'Check coordinates and try again.')                
        Tdistance += distance

        now = [Altstart,connect_point] 
        try: 
            distance=ors_distance(now)
        except:
           def Coorderror (request,id):
                messages.error(request, 'Check coordinates and try again.')                
        Tdistance += distance
    return Tdistance, Start, Altstart       

def Set_Start_Distance(points, set_start, connect_point=None):
    Start,Altstart= Gmaps_distance(points)
    for i in points: 
        Tdistance=0    
        now=[set_start,i]
        try:
            distance=ors_distance(now)
        except:
           def Coorderror (request,id):
                messages.error(request, 'Check coordinates and try again.') 
        Tdistance+=distance

    if connect_point:
        now=[set_start, connect_point]
        try:
            distance=ors_distance(now)
        except:
           def Coorderror (request,id):
                messages.error(request, 'Check coordinates and try again.') 
        Tdistance += distance

        now = [Altstart,connect_point]  
        try:
            distance=ors_distance(now)
        except:
           def Coorderror (request,id):
                messages.error(request, 'Check coordinates and try again.') 
        Tdistance += distance
    Start=set_start
   
    return Tdistance, Start, Altstart  


def plot(points, names, map_name, set_start=None, connect_point=None):
    if set_start:
        Tdistance, Start, Altstart= Set_Start_Distance(points, set_start, connect_point)
    else:
        Tdistance, Start, Altstart = Start_Distance(points, connect_point)

    
    gmaps.configure(api_key=apikey)

    fig=gmaps.figure()
    drawing = gmaps.drawing_layer()
    fig.add_layer(drawing)
    
    for i in range(len(points)):
        if Start == points[i]:
            continue
  
        now=[Start,points[i]]
        try:
            route=ors_route(now)
            geojson_layer = gmaps.geojson_layer(route)
            fig.add_layer(geojson_layer)
        except:
            def Coorderror (request,id):
             messages.error(request, 'Check coordinates and try again.')        


    location2=list()
    for i in points:
        start=i
        location2.append((start[1],start[0]))
    marker_layer = gmaps.marker_layer((location2), 
                                    hover_text='Client',                                         
                                    info_box_content=names)

    fig.add_layer(marker_layer)

    Start1=list()
    Start1.append((Start[1], Start[0]))
    marker_layer = gmaps.marker_layer((Start1), 
                                hover_text='Build Cabinet', 
                                label='OLT',
                                info_box_content='Network Start Point')
    fig.add_layer(marker_layer)
    
    if connect_point:
        now=[Start, connect_point]
        try:
            route = ors_route(now)    
            geojson_layer = gmaps.geojson_layer(route)
            fig.add_layer(geojson_layer)
        except:
           def Coorderror (request,id):
                messages.error(request, 'Check coordinates and try again.')                

        location3=list()
        location3.append((connect_point[1], connect_point[0]))
        marker_layer = gmaps.marker_layer((location3), 
                                        hover_text='Integartion Point',
                                        label='IP',
                                        info_box_content='Network Integration Point')
        fig.add_layer(marker_layer)

        now=list()
        now.append((connect_point[0],connect_point[1]))
        now.append((Altstart[0],Altstart[1]))
        try:
            route = ors_route(now)
            geojson_layer = gmaps.geojson_layer(route)
            fig.add_layer(geojson_layer)
        except:
            def Coorderror (request,id):
                messages.error(request, 'Check coordinates and try again.')    



    Altstart1=list()
    Altstart1.append((Altstart[1], Altstart[0]))
    marker_layer = gmaps.marker_layer((Altstart1), 
                            hover_text='Redundant Link Point', 
                            label='RP',
                            info_box_content='Redundant Start Point')
    fig.add_layer(marker_layer)


    embed_minimal_html(map_name, views=[fig])

    return Tdistance, Start


def item_count(plot):

    total_fibre_excess = math.ceil(plot.distance / plot.Extra_Fiber_Length_After) *  plot.Extra_Fiber_Length
    
    poles = math.ceil(plot.distance / plot.pole_separation) + 1 
    fibre_optic = plot.distance + total_fibre_excess
    man_holes = math.floor(plot.distance / plot.man_hole_separation) + 1
    hand_holes = math.floor(plot.distance / plot.hand_hole_separation)
    

    return poles, fibre_optic, man_holes, hand_holes


# Client = ors.Client(key='5b3ce3597851110001cf624820c44dd22f3c4a128773998371409a6f')
# # points=[[36.8159, -1.2795], [36.8219,-1.2921],[36.8259, -1.2850], [36.8145, -1.2870], [36.8222, -1.2935]]


# def ors_route(coordinates, output='map'):
#     route=Client.directions(coordinates=coordinates, 
#                             profile='foot-walking',
#                             optimized= False,
#                             preference='shortest',
#                             continue_straight=True,                        
#                             format='geojson')
#     if(output == 'map'):
#         return route
        
#     if(output == 'distance'):
#         return route['features'][0]['properties']['segments'][0]['distance']

# def find_start(points, connect_point=None):
#     MinTot=999999999
#     MaxTot = 0
#     Start=list()
#     for i in points: 
#         Tdistance=0
#         for y in points:
#             now=[i,y]
#             distance=ors_route(now, 'distance')
#             Tdistance+=distance
            

#         if MinTot>Tdistance:
#             MinTot=Tdistance
#             Start=i

#        #Altstart to be used to create redundancy for network security 
#         if MaxTot < Tdistance:
#             MaxTot = Tdistance
#             Altstart = i
      

#      # Adding the distance between the start and connect point
#     if connect_point:
#         now=[Start, connect_point]
#         distance=ors_route(now, 'distance')
#         MinTot += distance

#         now = [Altstart,connect_point]  
#         distance=ors_route(now, 'distance')
#         MinTot += distance
      

#     return MinTot, Start, Altstart


# def min_tot_for_set_point(set_start, points, connect_point=None):
#    Altstart=list()
#    MaxTot = 0
#    for i in points: 
#         Tdistance=0
#         for y in points:
#             now=[i,y]
#             distance=ors_route(now, 'distance')
#             Tdistance+=distance
#      #Altstart to be used to create redundancy for network security 
#         if MaxTot < Tdistance:
#             MaxTot = Tdistance
#             Altstart = i
    
#    if connect_point:
#         now=[set_start, connect_point]
#         distance=ors_route(now, 'distance')
#         Tdistance+=distance

#         now = [Altstart,connect_point]  
#         distance=ors_route(now, 'distance')
#         Tdistance += distance
        
#    Start = set_start   


#    return Tdistance, Start, Altstart


# def plot(points, names, map_name, Altstart, connect_point=None, set_start=None):
#     if set_start:
#        MinTot, Start, Altstart = min_tot_for_set_point(set_start, points, connect_point)
#     else:
#        MinTot, Start, Altstart = find_start(points, connect_point)

#     # Start=points[0]
#     my_directions=folium.Map(location=[Start[1], Start[0]], zoom_start=14, min_zoom=5, max_zoom= 16)  
#     for i in range(len(points)):
        
#         now=[Start,points[i]]
#         route = ors_route(now)
#         folium.GeoJson(route, name='route').add_to(my_directions)

#         folium.Marker(
#             location=[points[i][1], points[i][0]],
#             tooltip = names[i],
#             icon=folium.Icon(icon="home"),
#         ).add_to(my_directions)

   
    
#     if connect_point:
#         now=[Start, connect_point]
#         route = ors_route(now)
#         folium.GeoJson(route, name='route').add_to(my_directions)

#         folium.Marker(
#             location=[connect_point[1], connect_point[0]],
#             tooltip = 'Integration Point',
#             icon=folium.Icon(icon="info-sign", color='blue'),
#         ).add_to(my_directions)

#     now = [Altstart, connect_point]
#     route = ors_route(now)
#     folium.GeoJson(route, name='route').add_to(my_directions)


#     folium.Marker(
#               location=[Altstart[1], Altstart[0]],
#               tooltip = 'Redundant Link Point',
#               icon=folium.Icon(icon="down", color = 'green' ),
#              ).add_to(my_directions)

#     folium.Marker(
#             location=[Start[1], Start[0]],
#             tooltip = 'Build Cabinet',
#             icon=folium.Icon(icon="down", color='lightgreen'),
#         ).add_to(my_directions)



#     my_directions.save(map_name)
    

#     return MinTot, Start


#def item_count(plot):

#    total_fibre_excess = math.ceil(plot.distance / plot.Extra_Fiber_Length_After) *  plot.Extra_Fiber_Length
    
#    poles = math.ceil(plot.distance / plot.pole_separation) + 1 
 #   fibre_optic = plot.distance + total_fibre_excess
  #  man_holes = math.floor(plot.distance / plot.man_hole_separation) + 1
   # hand_holes = math.floor(plot.distance / plot.hand_hole_separation)
    

    #return poles, fibre_optic, man_holes, hand_holes

from docx import Document
from docx.shared import Inches

# https://python-docx.readthedocs.io/en/latest/#what-it-can-do
def save_quote_to_file(plot, poles, poles_price,
fibre_optic, fibre_optic_price,
man_holes, man_hole_price,
hand_holes, hand_hole_price,
Support_Tangent, support_tangent_price,
onu, onu_price,
olt, olt_price, 
total
):
    document = Document()

    # document.add_heading('Document Title', 0)

    # p = document.add_paragraph('A plain paragraph having some ')
    # p.add_run('bold').bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True

    # document.add_heading('Heading, level 1', level=1)
    # document.add_paragraph('Intense quote', style='Intense Quote')

    # document.add_paragraph(
    #     'first item in unordered list', style='List Bullet'
    # )
    # document.add_paragraph(
    #     'first item in ordered list', style='List Number'
    # )

    # document.add_picture('monty-truth.png', width=Inches(1.25))

    # records = (
    #     (3, '101', 'Spam'),
    #     (7, '422', 'Eggs'),
    #     (4, '631', 'Spam, spam, eggs, and spam')
    # )

    # table = document.add_table(rows=1, cols=3)
    # hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = 'Qty'
    # hdr_cells[1].text = 'Id'
    # hdr_cells[2].text = 'Desc'
    # for qty, id, desc in records:
    #     row_cells = table.add_row().cells
    #     row_cells[0].text = str(qty)
    #     row_cells[1].text = id
    #     row_cells[2].text = desc

    records = (
        (1, "Poles", plot.pole.sku, poles, plot.pole.unit_price, poles_price),
        (2, "Fibre Cable", plot.fibre_optic.sku, fibre_optic, plot.fibre_optic.unit_price, fibre_optic_price),
        (3, "Man hole", plot.man_hole.sku, man_holes, plot.man_hole.unit_price, man_hole_price),
        (4, "Hand Hole", plot.hand_hole.sku, hand_holes, plot.hand_hole.unit_price, hand_hole_price),
        (5, "Support Tangent", plot.Support_Tangent.sku, Support_Tangent, plot.Support_Tangent.unit_price, support_tangent_price),
        (6, "ONU", plot.onu.sku, onu, plot.onu.unit_price, onu_price),
        (7, "OLT", plot.olt.sku, olt, plot.olt.unit_price, olt_price),
        ("", "", "TOTAL", "", "", total),
    )

    table = document.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '#'
    hdr_cells[1].text = 'Category'
    hdr_cells[2].text = 'SKU'
    hdr_cells[3].text = 'Quantity'
    hdr_cells[4].text = 'Unit Price'
    hdr_cells[5].text = 'Price'
    
    for index, category, sku, quantity, unit_price, price in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(index)
        row_cells[1].text = category
        row_cells[2].text = sku
        row_cells[3].text = str(quantity)
        row_cells[4].text = str(unit_price)
        row_cells[5].text = str(price)

    document.add_page_break()

    # document.save('demo.docx')

    filename = f'media/{plot.name}.docx'
    # Thttps://www.w3schools.com/python/python_file_remove.asp
    import os
    if os.path.exists(filename):
        os.remove(filename)

    else:
        document.save(filename)